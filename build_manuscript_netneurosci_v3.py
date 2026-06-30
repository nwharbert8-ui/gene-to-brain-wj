# -*- coding: utf-8 -*-
"""
Network Neuroscience manuscript (v3) for the gene-to-brain cross-scale study.
Built from the 2026-05-30 rebuild + 2026-06-30 per-subject/three-state exploration.
Author: Drake H. Harbert (D.H.H.) | ORCID 0009-0007-7740-3616.
Conservative language, biological framing, honest nulls and limitations, no dashes.
"""
import os, sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
OUT = r"G:/My Drive/inner_architecture_research/gene_to_brain_wj/NetNeurosci_Submission_v3"
os.makedirs(OUT, exist_ok=True)
doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(1.0); s.left_margin = s.right_margin = Inches(1.0)
st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(12); st.paragraph_format.line_spacing = 2.0
C = WD_ALIGN_PARAGRAPH.CENTER
def P(t, bold=False, italic=False, align=None, after=0):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 2.0
    if align: p.alignment = align
    r = p.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(12); r.font.name = "Times New Roman"; return p
def H(t): return P(t, bold=True, after=2)

# TITLE PAGE
P("Gene co-expression architecture corresponds to human brain functional connectivity, "
  "and the correspondence is reversibly strengthened in the propofol-sedated state",
  bold=True, align=C, after=10)
P("Drake H. Harbert", align=C, after=2)
P("Inner Architecture LLC, Canton, OH 44720, United States", align=C, after=2)
P("ORCID: 0009-0007-7740-3616", align=C, after=2)
P("Correspondence: Drake@innerarchitecturellc.com", align=C, after=10)
P("Running title: Gene co-expression and brain connectivity", italic=True)
doc.add_page_break()

# ABSTRACT (<=200 words, unstructured)
H("Abstract")
P("Whether the molecular architecture of the brain constrains its functional "
  "organization across spatial scales remains open. We asked whether the pairwise "
  "co-expression architecture of genes across human brain regions corresponds to the "
  "functional connectivity architecture of those regions, and whether the "
  "correspondence depends on conscious state. Using genome-wide co-expression across "
  "eleven brain regions (GTEx) and resting and task functional connectivity in 24 "
  "individuals before, during, and after propofol sedation, we "
  "quantified co-expression architecture with a weighted Jaccard index and tested its "
  "correspondence to connectivity with Mantel permutation (respecting region-pair "
  "non-independence). Per individual, co-expression architecture "
  "corresponded to connectivity in all three states, and the correspondence was "
  "reversibly strengthened in the propofol-sedated state (within-subject increase of "
  "0.16, p = 6e-05; recovery returned to baseline). The strengthening survived "
  "adjustment for head motion and was independent of inter-regional distance and "
  "cell-type composition, and held within anatomical blocks. "
  "Because the sedated condition differed in task as well as drug, the modulation "
  "reflects the sedated state rather than propofol specifically. Gene expression level "
  "alone showed no such correspondence. These findings indicate that molecular "
  "co-expression architecture is associated with a relatively stable organizational "
  "component of functional connectivity that becomes more apparent when activity is "
  "suppressed.",
  after=6)
P("Keywords: gene co-expression, functional connectivity, cross-scale, weighted Jaccard "
  "index, propofol, consciousness, Mantel test", after=0)
doc.add_page_break()

# INTRODUCTION
H("1. Introduction")
for t in [
 "Brain function is organized across spatial scales, from molecular interactions within "
 "cells to coordinated activity among distributed regions. A long-standing question is "
 "whether the molecular architecture of the brain, the pattern of how genes are "
 "co-expressed across regions, is reflected in the macroscale functional organization "
 "measured by neuroimaging. Regional gene expression has been related to functional and "
 "structural connectivity in several studies, but most analyses correlate expression "
 "profiles or individual genes with connectivity, rather than comparing the relational "
 "architecture of co-expression to the relational architecture of connectivity.",
 "Here we treat both scales as architectures. At the molecular scale, the architecture "
 "is the full set of pairwise co-expression relationships among genes within each brain "
 "region; the similarity of these relationship patterns between two regions is "
 "quantified with a weighted Jaccard index (Harbert, 2026). At the macroscale, the "
 "architecture is the functional connectivity among the same regions. We ask whether "
 "the between-region similarity of co-expression architecture corresponds to the "
 "between-region functional connectivity, and whether this correspondence is fixed or "
 "depends on the brain's state.",
 "Propofol-induced unconsciousness provides a within-subject manipulation of state that "
 "suppresses activity-dependent dynamics while preserving the anatomical substrate. If "
 "molecular architecture is associated with a relatively stable component of connectivity, the "
 "correspondence might be obscured by activity-dependent variation in the awake state "
 "and become more apparent when that variation is reduced. We therefore measured the "
 "correspondence in each individual across the awake, unconscious, and recovery states, "
 "using an inferential framework appropriate to the non-independence of region pairs.",
]: P(t, after=4)
doc.add_page_break()

# METHODS
H("2. Materials and Methods")
H("2.1 Gene co-expression data")
P("Genome-wide expression (GTEx v8) was used to compute, within each of eleven brain "
  "regions, the full pairwise Spearman correlation matrix among 16,273 expressed genes "
  "(2,268 samples total; individual samples as the unit, no pre-aggregation). Regions "
  "were retained if at least 30 samples were available.", after=4)
H("2.2 Functional connectivity data")
P("Functional connectivity was taken from a publicly available propofol fMRI dataset "
  "(OpenNeuro ds006623; 24 individuals with usable data in all states). Data were "
  "preprocessed with XCP-D (motion regression, without global signal regression) and "
  "parcellated (4S456 atlas). Three states were defined as in the source acquisition: "
  "awake (resting, pre-propofol), unconscious (post loss of response, propofol "
  "maintenance, acquired during a mental-imagery task), and recovery (resting, "
  "post-propofol) (Huang et al., 2026). Behavioral unresponsiveness under sedation "
  "does not by itself establish unconsciousness; this dataset was designed to probe "
  "covert consciousness. Parcel time series were "
  "aggregated to the eleven gene-expression regions, and a Spearman connectivity matrix "
  "was computed per individual per state, and as a group average.", after=4)
H("2.3 Architecture comparison")
P("For each pair of regions, co-expression architecture similarity was the unsigned "
  "weighted Jaccard index of their gene-gene correlation matrices (the sum over gene "
  "pairs of the minimum of the two absolute correlations divided by the sum of the "
  "maxima). Correspondence between the co-expression architecture and the functional "
  "connectivity architecture was tested with the Mantel permutation test (10,000 "
  "permutations of whole regions, preserving the non-independence of matrix-derived "
  "pairs). Inter-regional Euclidean distance was controlled with partial Mantel. "
  "Expression-profile similarity (Spearman of regional median expression) was tested as "
  "a comparison.", after=4)
P("The weighted Jaccard index was chosen empirically and on construction. Empirically, "
  "among candidate architecture-similarity metrics applied to the same matrices it "
  "yielded the strongest correspondence to functional connectivity (Supplementary Table "
  "S2: weighted Jaccard |r| = 0.61, cosine 0.59, Frobenius distance 0.57, Pearson matrix "
  "correlation 0.45, Spearman matrix correlation 0.41, expression-profile similarity "
  "0.32), clearly exceeding standard matrix correlation, distance-based comparison, and "
  "expression level. By construction it is the bounded continuous generalization of the "
  "Jaccard index (shared architecture over total architecture), the construct of interest "
  "here; unlike the RV coefficient it is not dominated by high-variance dimensions, and "
  "unlike Procrustes analysis it requires no alignment or shared dimensionality. The "
  "Mantel comparison step is a representational similarity analysis; the weighted Jaccard "
  "index is the choice of how the molecular-side matrix is built.", after=4)
H("2.4 Per-subject and state analyses")
P("For each individual and state, the Mantel correspondence between the group "
  "co-expression architecture and that individual's connectivity was computed. The "
  "distribution of per-subject correspondences was tested against zero (Wilcoxon "
  "signed-rank), with individuals as the independent unit. Within-subject state "
  "differences (the state-treatment pairing) were tested paired, with bootstrap "
  "confidence intervals (1,000 iterations). Per-region contributions were computed as "
  "each region's row correspondence (its co-expression coupling profile versus its "
  "connectivity profile) per individual per state (Layer 2F decomposition).", after=4)
H("2.5 Cell-type and motion control")
P("Cell-type composition was addressed by estimating per-region scores for five "
  "canonical neural cell populations (neuron, astrocyte, oligodendrocyte, microglia, "
  "endothelial) from marker genes, forming a region-by-region cell-type composition "
  "distance, and recomputing the correspondence with partial Mantel controlling for it. "
  "Head motion was assessed in all three states (mean framewise displacement per "
  "individual); the per-subject change in correspondence was regressed on the per-subject "
  "change in framewise displacement, the motion-adjusted state effect being the "
  "intercept. Dependence on the subcortical-versus-cortical block was tested by partial "
  "Mantel controlling a same-block indicator and by restricting to within-block pairs. "
  "Random seed 42 was used throughout. Analysis code is available (Data and Code "
  "Availability).", after=4)
doc.add_page_break()

# RESULTS
H("3. Results")
H("3.1 Co-expression architecture corresponds to functional connectivity")
P("At the group level, co-expression architecture similarity corresponded to functional "
  "connectivity under unconsciousness (Mantel r = 0.610, p = 0.003) and at recovery "
  "(r = 0.458, p = 0.051), but the awake-state correspondence did not reach significance "
  "(r = 0.392, p = 0.11). The unconscious correspondence survived control for "
  "inter-regional distance (partial Spearman rho = 0.297, p = 0.027). Expression-profile "
  "similarity did not correspond to connectivity in any state (awake r = 0.04, p = 0.91).", after=4)
H("3.2 The correspondence is present in every individual and in all three states")
P("Computed per individual, with individuals as the independent unit (n = 24), the "
  "correspondence was significant in all three states: awake mean Mantel r = 0.338 (96% "
  "of individuals positive), unconscious 0.498 (100% positive), recovery 0.358 (100% "
  "positive); each distribution differed from zero (Wilcoxon p < 0.001). The single "
  "group matrix gave larger point estimates but weaker inference; the per-subject "
  "analysis is the stronger footing and treats individuals, not non-independent region "
  "pairs, as the unit.", after=4)
H("3.3 The correspondence is reversibly strengthened in the propofol-sedated state")
P("Within individuals, the correspondence was higher during unconsciousness than awake "
  "(paired difference 0.160, 95% CI 0.100 to 0.219, p = 6e-05) and than recovery "
  "(difference 0.140, p = 0.0007), while recovery and awake did not differ (difference "
  "0.019, p = 0.53). The unconscious state was the within-subject maximum in 71% of "
  "individuals, and the strengthening reversed on recovery. The correspondence is thus "
  "modulated by state in a reversible manner. The unconscious condition was acquired "
  "during an imagery task while the awake and recovery conditions were resting "
  "(Limitations), so the modulation is attributable to the unconscious state as a whole "
  "rather than to propofol specifically. Per-individual correspondence did not track "
  "propofol effect-site concentration at loss of response in any state.", after=4)
H("3.4 The correspondence is not explained by cell-type composition or distance")
P("Cell-type composition was associated with co-expression architecture (Mantel "
  "r = -0.71, p = 0.002) but not with functional connectivity (r = -0.24, p = 0.36). "
  "The unconscious correspondence survived control for cell-type composition (partial "
  "r = 0.643, p = 3e-04) and for distance and cell-type composition together (r = 0.428, "
  "p = 0.005). The part of co-expression architecture that corresponds to connectivity "
  "is therefore independent of cell-type composition and of physical distance.", after=4)
H("3.5 Anatomical localization")
P("Per-region decomposition showed that the correspondence was carried most strongly by "
  "the basal ganglia (putamen, nucleus accumbens, caudate, substantia nigra; row "
  "correspondences 0.53 to 0.77 across states), where it was high in every state. The "
  "state-dependent strengthening was concentrated in cortex and cerebellum: in the awake "
  "state these regions showed low or negative correspondence (frontal cortex -0.48, "
  "cerebellum -0.52), which rose toward the molecular pattern under unconsciousness "
  "(frontal cortex within-subject change 0.45, cerebellum 0.41, both p < 0.001). The "
  "hippocampus was the only region in which correspondence decreased under "
  "unconsciousness (change -0.10, uncorrected p = 0.037); this did not survive "
  "correction across regions and is not interpreted further.", after=4)
H("3.6 Head motion does not account for the state effect")
P("Framewise displacement was higher in the sedated state (means 0.147 awake, 0.211 "
  "sedated, 0.163 recovery). To test whether motion produces the strengthening, the "
  "per-subject change in correspondence (sedated minus awake) was regressed on the "
  "per-subject change in framewise displacement. The motion-adjusted state effect, the "
  "intercept, remained positive and significant (0.179, p = 0.001), and the slope on "
  "framewise displacement was not significant (p = 0.60). Consistent with this, the "
  "change in correspondence did not correlate with the change in motion across "
  "individuals (Spearman = -0.085, p = 0.69). The positive motion-adjusted effect is the "
  "primary evidence; the two null associations corroborate it but, at this sample size, "
  "are not on their own decisive. Head motion does not account for the state effect.", after=4)
H("3.7 The correspondence is not the subcortical-cortical dichotomy")
P("To test whether the correspondence merely reflects subcortical and cortical regions "
  "clustering by type, it was recomputed controlling for a same-block indicator. The "
  "block-controlled correspondence was essentially unchanged (partial Mantel r = 0.56, "
  "p = 0.003) and was robust across block definitions and across Pearson and Spearman "
  "estimation. Restricting to within-block pairs preserved a positive correspondence "
  "(illustratively, within the subcortical block r = 0.67 on a small number of pairs; the "
  "cortical block has too few regions to estimate reliably). The correspondence is "
  "therefore architecture within blocks, not only the coarse dichotomy between them.", after=4)
doc.add_page_break()

# DISCUSSION
H("4. Discussion")
for t in [
 "Across 24 individuals, the pairwise co-expression architecture of genes among human "
 "brain regions corresponds to the functional connectivity architecture of those "
 "regions. The correspondence is present in each individual and in every state, is not "
 "explained by inter-regional distance or by cell-type composition, and is not produced "
 "by expression level alone. Its strength is modulated by conscious state in a "
 "reversible manner: it is higher in the sedated state and returns to baseline on "
 "recovery. The strengthening is not accounted for by head motion and is not the coarse "
 "subcortical-cortical dichotomy.",
 "The pattern is consistent with co-expression architecture being associated with a "
 "relatively stable organizational component of functional connectivity that is partly "
 "obscured by activity-dependent variation in the awake brain and becomes more apparent "
 "when that variation is "
 "suppressed. The anatomical localization supports this reading: the basal ganglia "
 "express the correspondence in every state, whereas cortex and cerebellum, where "
 "awake connectivity diverges most from the molecular pattern, move toward it under "
 "unconsciousness. We frame this as an association rather than a causal or directional "
 "claim; the analysis is correlational.",
 "Methodologically, treating both scales as architectures and comparing them with a "
 "permutation test that respects the non-independence of region pairs avoids the "
 "inflation that arises when matrix-derived pairs are treated as independent "
 "observations. The per-subject formulation, with individuals as the unit, provides the "
 "primary inference and shows that the correspondence is a property of individuals, not "
 "an artifact of a single averaged matrix.",
]: P(t, after=4)
H("4.1 Limitations")
P("The analysis is correlational and does not establish that molecular architecture "
  "causes or directs functional connectivity. The gene-expression and connectivity data "
  "come from different individuals (regional architecture from GTEx, connectivity from "
  "the propofol cohort), so the correspondence is at the level of regional architecture, "
  "not within-person molecular-to-functional mapping. The unconscious condition was "
  "acquired during a mental-imagery, covert-consciousness task while awake and recovery "
  "were resting, so sedation and task are confounded for that condition, and behavioral "
  "unresponsiveness does not establish unconsciousness (Huang et al., 2026). Framewise "
  "displacement was higher in the sedated state (0.15 awake versus 0.21 sedated); "
  "however, the state effect survives adjustment for framewise displacement and does not "
  "track motion across individuals (Section 3.6), and motion artifact in functional "
  "connectivity is predominantly distance dependent (Power et al., 2012), which the "
  "distance-controlled analysis also addresses, so motion does not account for the "
  "finding. Eleven regions limit spatial resolution, and per-region estimates rest on few "
  "partners and are reported only at the group level. Surface-based "
  "spatial-autocorrelation null models were not applied at this regional resolution; "
  "distance control and distance-matched subsampling were used instead. Finally, the data "
  "are from a single propofol cohort; replication in an independent cohort and a different "
  "anesthetic (for example sevoflurane or ketamine, or the Cambridge propofol dataset) "
  "would test whether the strengthening is specific to propofol or general to suppressed "
  "states, which the task confound here leaves open.", after=4)
H("4.2 Conclusions")
P("Molecular co-expression architecture corresponds to the functional connectivity "
  "architecture of the human brain, in every individual examined and across conscious "
  "states, and the correspondence is reversibly strengthened in the sedated state, an "
  "effect not accounted for by head motion, distance, cell-type composition, or the "
  "subcortical-cortical dichotomy. The result is consistent with molecular architecture "
  "being associated with a relatively stable organizational component of connectivity "
  "that becomes more apparent as activity-dependent dynamics recede.", after=4)
doc.add_page_break()

# DECLARATIONS
H("Declaration of Generative AI Use")
P("Claude (Anthropic, Claude Opus 4.8) was used as a programming assistant during "
  "pipeline development, manuscript formatting, and code review. All analytical "
  "decisions, methodology design, data interpretation, and scientific conclusions are "
  "solely the work of the author. The AI tool was not used to generate scientific text, "
  "interpret results, or formulate hypotheses. All code was reviewed and validated by "
  "the author prior to execution.", after=4)
H("Data and Code Availability")
P("Gene expression data are from GTEx v8 (gtexportal.org). Functional connectivity data "
  "are from OpenNeuro dataset ds006623 (Huang et al., 2026). All analysis code that "
  "reproduces the results is openly available at "
  "https://github.com/nwharbert8-ui/gene-to-brain-wj.", after=4)
H("Author Contributions (CRediT)")
P("Drake H. Harbert: Conceptualization, Methodology, Software, Formal Analysis, "
  "Investigation, Data Curation, Writing - Original Draft, Writing - Review and Editing, "
  "Visualization.", after=4)
H("Competing Interests")
P("The author is the founder of Inner Architecture LLC and declares no other competing "
  "interests.", after=4)
H("Ethics")
P("This study used publicly available, de-identified data (GTEx; OpenNeuro ds006623). "
  "All original data were collected with institutional review board approval at the "
  "originating institutions. No additional approval was required for secondary analysis.", after=4)
H("Funding")
P("This research received no external funding.", after=4)
doc.add_page_break()

# REFERENCES (APA; all must resolve; NO in-press per MIT policy)
H("References")
refs = [
 "GTEx Consortium. (2020). The GTEx Consortium atlas of genetic regulatory effects "
 "across human tissues. Science, 369(6509), 1318-1330. https://doi.org/10.1126/science.aaz1776",
 "Harbert, D. H. (2026). Sigma-1 and sigma-2 receptors exhibit divergent genome-wide "
 "co-expression architectures in human brain despite shared subcellular localization. "
 "Frontiers in Pharmacology, 17, 1830847. https://doi.org/10.3389/fphar.2026.1830847",
 "Langfelder, P., & Horvath, S. (2008). WGCNA: an R package for weighted correlation "
 "network analysis. BMC Bioinformatics, 9, 559. https://doi.org/10.1186/1471-2105-9-559",
 "Mantel, N. (1967). The detection of disease clustering and a generalized regression "
 "approach. Cancer Research, 27(2), 209-220.",
 "Power, J. D., Barnes, K. A., Snyder, A. Z., Schlaggar, B. L., & Petersen, S. E. "
 "(2012). Spurious but systematic correlations in functional connectivity MRI networks "
 "arise from subject motion. NeuroImage, 59(3), 2142-2154. "
 "https://doi.org/10.1016/j.neuroimage.2011.10.018",
 "Huang, Z., Tarnal, V., Fotiadis, P., Vlisides, P. E., Janke, E. L., Puglia, M., "
 "McKinney, A. M., Jang, H., Dai, R., Picton, P., Mashour, G. A., & Hudetz, A. G. "
 "(2026). An open fMRI resource for studying human brain function and covert "
 "consciousness under anesthesia. Scientific Data, 13, 127. "
 "https://doi.org/10.1038/s41597-025-06442-2",
]
for r in refs:
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 2.0; p.paragraph_format.space_after = Pt(2)
    rr = p.add_run(r); rr.font.name = "Times New Roman"; rr.font.size = Pt(12)

doc.save(os.path.join(OUT, "Manuscript.docx"))
print("WROTE:", os.path.join(OUT, "Manuscript.docx"), "| words:",
      sum(len(p.text.split()) for p in doc.paragraphs))
