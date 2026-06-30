# -*- coding: utf-8 -*-
"""
Figures and main table for the gene-to-brain Network Neuroscience manuscript.
Author: Drake H. Harbert (D.H.H.) | ORCID 0009-0007-7740-3616 | 2026-06-30
Reads the analysis outputs and renders 300 DPI colorblind-safe figures + Table 1.
"""
import os, sys, json
import numpy as np, pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
B = r"G:/My Drive/inner_architecture_research/gene_to_brain_wj/results"
OUT = r"G:/My Drive/inner_architecture_research/gene_to_brain_wj/NetNeurosci_Submission_v3"
FIG = os.path.join(OUT, "Main_Figures"); TAB = os.path.join(OUT, "Main_Tables")
os.makedirs(FIG, exist_ok=True); os.makedirs(TAB, exist_ok=True)
CB = {"awake": "#0072B2", "unconscious": "#D55E00", "recovery": "#009E73"}
CONDS = ["awake", "unconscious", "recovery"]

# ---------- FIGURE 1: per-subject correspondence across 3 states ----------
ps = pd.read_csv(os.path.join(B, "per_subject", "per_subject_mantel.csv"))
fig, ax = plt.subplots(figsize=(8, 6))
data = [ps[f"mantel_r_{c}"].dropna().values for c in CONDS]
positions = [1, 2, 3]
for i, c in enumerate(CONDS):
    ax.scatter(np.random.RandomState(42 + i).normal(positions[i], 0.05, len(data[i])), data[i],
               color=CB[c], alpha=0.6, s=40, zorder=3, edgecolors="black", linewidths=0.4)
# paired lines per subject
for _, row in ps.iterrows():
    ys = [row[f"mantel_r_{c}"] for c in CONDS]
    ax.plot(positions, ys, color="#999999", alpha=0.35, lw=0.7, zorder=1)
bp = ax.boxplot(data, positions=positions, widths=0.5, showfliers=False, zorder=2,
                medianprops=dict(color="black", lw=2), boxprops=dict(color="black"))
ax.axhline(0, color="black", lw=0.8, ls="--", alpha=0.6)
ax.set_xticks(positions); ax.set_xticklabels(["Awake", "Unconscious", "Recovery"], fontsize=12)
ax.set_ylabel("Per-subject gene-to-FC correspondence (Mantel r)", fontsize=12)
ax.set_title("Co-expression to connectivity correspondence per individual, by state\n"
             "(reversibly strengthened during unconsciousness; n = 24)", fontsize=12)
ax.text(0.5, 0.02, "unconscious - awake = +0.16, p = 6e-05; recovery = awake (p = 0.53)",
        transform=ax.transAxes, ha="center", fontsize=9, style="italic")
plt.tight_layout(); plt.savefig(os.path.join(FIG, "Fig1.png"), dpi=300); plt.savefig(os.path.join(FIG, "Fig1.pdf")); plt.close()
print("Fig1 done")

# ---------- FIGURE 2: per-region row correspondence by state ----------
pr = pd.read_csv(os.path.join(B, "per_region_state", "per_region_state.csv"))
pr = pr.sort_values("mean_across_states", ascending=False)
labels = [r.replace("Brain - ", "").replace(" (basal ganglia)", "") for r in pr["region"]]
x = np.arange(len(labels)); w = 0.26
fig, ax = plt.subplots(figsize=(11, 6))
for k, c in enumerate(CONDS):
    ax.bar(x + (k - 1) * w, pr[f"rowcorr_{c}"], w, label=c.capitalize(), color=CB[c], edgecolor="black", linewidth=0.4)
ax.axhline(0, color="black", lw=0.8)
ax.set_xticks(x); ax.set_xticklabels(labels, rotation=40, ha="right", fontsize=10)
ax.set_ylabel("Per-region row correspondence (Pearson)", fontsize=12)
ax.set_title("Regional contribution to the gene-to-FC correspondence, by state\n"
             "(basal ganglia carry it in all states; cortex and cerebellum are state-driven)", fontsize=12)
ax.legend(fontsize=11); plt.tight_layout()
plt.savefig(os.path.join(FIG, "Fig2.png"), dpi=300); plt.savefig(os.path.join(FIG, "Fig2.pdf")); plt.close()
print("Fig2 done")

# ---------- FIGURE 3: robustness controls (unconscious correspondence) ----------
ct = json.load(open(os.path.join(B, "celltype_confound", "celltype_confound.json")))["results"]
bars = [("WJ vs FC\n(unadjusted)", ct["WJ_vs_FCuncon"]["mantel_r"], ct["WJ_vs_FCuncon"]["p"]),
        ("partial:\ncell-type", ct["WJ_vs_FCuncon_PARTIAL_celltype"]["partial_mantel_r"], ct["WJ_vs_FCuncon_PARTIAL_celltype"]["p"]),
        ("partial:\ndistance+cell-type", ct["WJ_vs_FCuncon_PARTIAL_distance_AND_celltype"]["partial_r"], ct["WJ_vs_FCuncon_PARTIAL_distance_AND_celltype"]["p"])]
fig, ax = plt.subplots(figsize=(7, 6))
xs = np.arange(len(bars))
ax.bar(xs, [b[1] for b in bars], color="#D55E00", edgecolor="black", width=0.55)
for i, b in enumerate(bars):
    ax.text(i, b[1] + 0.02, f"r={b[1]:.2f}\np={b[2]:.0e}" if b[2] < 0.01 else f"r={b[1]:.2f}\np={b[2]:.3f}",
            ha="center", fontsize=9)
ax.set_xticks(xs); ax.set_xticklabels([b[0] for b in bars], fontsize=11)
ax.set_ylabel("Unconscious-state correspondence (Mantel r)", fontsize=12); ax.set_ylim(0, 0.8)
ax.set_title("The correspondence survives control for distance and cell-type composition", fontsize=12)
plt.tight_layout(); plt.savefig(os.path.join(FIG, "Fig3.png"), dpi=300); plt.savefig(os.path.join(FIG, "Fig3.pdf")); plt.close()
print("Fig3 done")

# ---------- TABLE 1: main results ----------
sp = json.load(open(os.path.join(B, "state_pairings", "state_pairings.json")))
doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.8); s.left_margin = s.right_margin = Inches(0.8)
doc.styles["Normal"].font.name = "Times New Roman"; doc.styles["Normal"].font.size = Pt(10)
p = doc.add_paragraph(); r = p.add_run("Table 1. Gene co-expression to functional connectivity correspondence."); r.bold = True; r.font.size = Pt(11)
p = doc.add_paragraph(); r = p.add_run("Per-subject Mantel correspondence (n = 24) and within-subject state differences; "
    "group single-matrix Mantel shown for reference. Inference is over independent subjects (Wilcoxon)."); r.font.size = Pt(9)
rows = [["Measure", "Awake", "Unconscious", "Recovery"]]
psum = sp.get("per_subject_vs_group", {})
rows.append(["Per-subject mean Mantel r",
             f"{psum['awake']['per_subject_mean']}", f"{psum['unconscious']['per_subject_mean']}", f"{psum['recovery']['per_subject_mean']}"])
rows.append(["Per-subject 95% CI",
             f"{psum['awake']['per_subject_ci']}", f"{psum['unconscious']['per_subject_ci']}", f"{psum['recovery']['per_subject_ci']}"])
rows.append(["Group single-matrix r (p)",
             f"{psum['awake']['group_single_matrix']} ({psum['awake']['group_p']})",
             f"{psum['unconscious']['group_single_matrix']} ({psum['unconscious']['group_p']})",
             f"{psum['recovery']['group_single_matrix']} ({psum['recovery']['group_p']})"])
st = sp["state_treatment_pairing"]
rows.append(["Within-subject state contrast", "", "", ""])
rows.append(["  unconscious - awake", f"{st['unconscious_minus_awake']['gap_mean']} (p={st['unconscious_minus_awake']['wilcoxon_p']})", "", ""])
rows.append(["  unconscious - recovery", f"{st['unconscious_minus_recovery']['gap_mean']} (p={st['unconscious_minus_recovery']['wilcoxon_p']})", "", ""])
rows.append(["  recovery - awake", f"{st['recovery_minus_awake']['gap_mean']} (p={st['recovery_minus_awake']['wilcoxon_p']})", "", ""])
t = doc.add_table(rows=len(rows), cols=4); t.style = "Light Grid Accent 1"
for i, rr in enumerate(rows):
    for j, v in enumerate(rr):
        c = t.rows[i].cells[j]; c.text = ""; run = c.paragraphs[0].add_run(str(v)); run.font.size = Pt(9); run.bold = (i == 0)
doc.save(os.path.join(TAB, "Table1.docx"))
print("Table1 done")

# ---------- TABLE S2: architecture-similarity metric comparison (why WJ) ----------
mc = json.load(open(os.path.join(B, "metric_comparison", "metric_comparison.json")))
d2 = Document()
for s in d2.sections:
    s.top_margin = s.bottom_margin = Inches(0.8); s.left_margin = s.right_margin = Inches(0.8)
d2.styles["Normal"].font.name = "Times New Roman"; d2.styles["Normal"].font.size = Pt(10)
p = d2.add_paragraph(); r = p.add_run("Supplementary Table S2. Architecture-similarity metric comparison."); r.bold = True; r.font.size = Pt(11)
p = d2.add_paragraph(); r = p.add_run("Correspondence of each candidate inter-regional architecture-similarity "
    "metric to unconscious-state functional connectivity (|Mantel r|, Pearson-based, 10,000 permutations; and "
    "|Spearman| of the upper-triangle vectors). Weighted Jaccard yields the strongest correspondence among "
    "tested metrics, clearly exceeding standard matrix correlation, distance-based comparison, and expression "
    "level. Frobenius is a distance (reported as |r|)."); r.font.size = Pt(9)
rows = [["Metric", "|Mantel r|", "p", "|Spearman|"]] + [[x["metric"], f"{x['mantel_r']:.3f}", f"{x['mantel_p']:.4f}", f"{x['spearman']:.3f}"] for x in mc["rows"]]
t = d2.add_table(rows=len(rows), cols=4); t.style = "Light Grid Accent 1"
for i, rr in enumerate(rows):
    for j, v in enumerate(rr):
        c = t.rows[i].cells[j]; c.text = ""; run = c.paragraphs[0].add_run(str(v)); run.font.size = Pt(9); run.bold = (i == 0)
d2.save(os.path.join(TAB, "TableS2_metric_comparison.docx"))
print("TableS2 done")
print("ALL FIGURES + TABLES WRITTEN to", OUT)
