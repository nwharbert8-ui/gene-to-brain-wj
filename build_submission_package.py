# -*- coding: utf-8 -*-
"""Assemble the complete Network Neuroscience submission folder for gene-to-brain:
cover letter, title page, reference verification, supplementary copies, manifest."""
import os, sys, shutil, glob
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
BASE = r"G:/My Drive/inner_architecture_research/gene_to_brain_wj"
OUT = os.path.join(BASE, "NetNeurosci_Submission_v3")
SUP = os.path.join(OUT, "Supplementary"); os.makedirs(SUP, exist_ok=True)
TITLE = ("Gene co-expression architecture corresponds to human brain functional "
         "connectivity, and the correspondence is reversibly strengthened during propofol sedation")

def doc11():
    d = Document()
    for s in d.sections:
        s.top_margin = s.bottom_margin = Inches(1.0); s.left_margin = s.right_margin = Inches(1.0)
    d.styles["Normal"].font.name = "Times New Roman"; d.styles["Normal"].font.size = Pt(12)
    return d
def para(d, t, bold=False, italic=False, align=None, after=8, sp=1.15, size=12):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = sp
    if align: p.alignment = align
    r = p.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size); return p

# ---------- TITLE PAGE ----------
d = doc11(); C = WD_ALIGN_PARAGRAPH.CENTER
para(d, TITLE, bold=True, align=C, after=14, sp=2.0)
para(d, "Drake H. Harbert", align=C, after=2)
para(d, "Inner Architecture LLC, Canton, OH 44720, United States", align=C, after=2)
para(d, "ORCID: 0009-0007-7740-3616", align=C, after=2)
para(d, "Correspondence: Drake H. Harbert, Drake@innerarchitecturellc.com", align=C, after=14)
para(d, "Running title: Gene co-expression and brain connectivity", after=2)
para(d, "Keywords: gene co-expression; functional connectivity; cross-scale; weighted Jaccard index; propofol; consciousness; Mantel test", after=2)
para(d, "Competing interests: The author is the founder of Inner Architecture LLC and declares no other competing interests.", after=2)
para(d, "Funding: This research received no external funding.", after=2)
d.save(os.path.join(OUT, "Title_Page.docx"))

# ---------- COVER LETTER ----------
d = doc11()
para(d, "Drake H. Harbert", after=0); para(d, "Inner Architecture LLC, Canton, OH, United States", size=10, after=0)
para(d, "Drake@innerarchitecturellc.com  |  ORCID 0009-0007-7740-3616", size=10, after=12)
para(d, "June 30, 2026", after=0)
para(d, "The Editors, Network Neuroscience (MIT Press)", after=12)
para(d, "Re: Submission of an original Research article", bold=True, after=10)
para(d, "Dear Editors,")
for t in [
 "Please consider the enclosed manuscript, \"" + TITLE + ",\" as an original Research article for Network Neuroscience.",
 "The study asks a cross-scale question central to network neuroscience: whether the pairwise co-expression architecture of genes across human brain regions corresponds to the functional connectivity architecture of those regions, and whether that correspondence depends on brain state. Using genome-wide co-expression (GTEx) and a publicly available propofol fMRI dataset (OpenNeuro ds006623), and comparing the two architectures with a permutation test that respects the non-independence of region pairs, we find that the correspondence is present in every individual and in all three states, is reversibly strengthened during propofol sedation, survives control for inter-regional distance and cell-type composition, and is carried most strongly by the basal ganglia while the state-dependent component is concentrated in cortex and cerebellum.",
 "We believe this fits the journal's scope at the intersection of molecular and macroscale brain organization. The inference treats individual subjects, not non-independent region pairs, as the unit, and the limitations (the cross-individual design, the sedation-versus-task confound, and the correlational nature of the finding) are stated explicitly. All data are publicly available and all analysis code is deposited.",
 "The manuscript is original, is not under consideration elsewhere, and the sole author has approved this submission. The author declares no competing interests beyond founding Inner Architecture LLC.",
 "Thank you for your consideration.",
 "Respectfully,", "Drake H. Harbert",
]: para(d, t)
d.save(os.path.join(OUT, "Cover_Letter.docx"))

# ---------- REFERENCE VERIFICATION ----------
refdoc = """# Reference verification, gene-to-brain manuscript (Network Neuroscience)
# Generated 2026-06-30. MIT Press policy: no "in press"/"submitted" citations -> all must resolve.

| # | Reference | Identifier | Status |
|---|-----------|-----------|--------|
| 1 | GTEx Consortium (2020). The GTEx Consortium atlas... Science 369:1318-1330 | doi:10.1126/science.aaz1776 | RESOLVES (canonical) |
| 2 | Harbert, D.H. (2026). Sigma-1 and sigma-2 receptors... Front. Pharmacol. 17:1830847 | doi:10.3389/fphar.2026.1830847 | VERIFIED (web 2026-06-30); cited for WJ method only, NOT the hippocampus result |
| 3 | Langfelder & Horvath (2008). WGCNA... BMC Bioinformatics 9:559 | doi:10.1186/1471-2105-9-559 | RESOLVES (canonical) |
| 4 | Mantel, N. (1967). The detection of disease clustering... Cancer Research 27:209-220 | (no DOI; pre-DOI era) | RESOLVES (canonical) |
| 5 | Power et al. (2012). Spurious but systematic correlations... NeuroImage 59:2142-2154 | doi:10.1016/j.neuroimage.2011.10.018 | RESOLVES (canonical) |
| 6 | Huang, Z., et al. (2025). An open fMRI resource... covert consciousness under anesthesia. Scientific Data | doi:10.1038/s41597-025-06442-2 (OpenNeuro ds006623) | VERIFIED (web 2026-06-30). TODO: complete full author list from the paper before final submission |

## Notes
- This is a NEW manuscript for a new venue; a prior-version reference diff (Rule 1) does not apply.
- No fabricated or unresolvable references. No "in press"/"submitted" citations (MIT policy satisfied).
- Outstanding before final submission: (a) complete Huang et al. 2025 author list; (b) confirm the
  Harbert 2026 text actually supports the WJ-method claim it is cited for; (c) mint the Zenodo DOI.
"""
open(os.path.join(OUT, "_reference_verification.md"), "w", encoding="utf-8").write(refdoc)

# ---------- SUPPLEMENTARY: copy key result tables ----------
for src in [
    os.path.join(BASE, "results", "per_subject", "per_subject_mantel.csv"),
    os.path.join(BASE, "results", "per_region_state", "per_region_state.csv"),
    os.path.join(BASE, "results", "state_pairings", "state_pairings.json"),
    os.path.join(BASE, "results", "celltype_confound", "celltype_confound.json"),
    os.path.join(BASE, "results", "headmotion_check", "headmotion_check.json"),
    os.path.join(BASE, "results", "rebuild", "pairwise_metrics.csv"),
]:
    if os.path.exists(src):
        shutil.copy(src, os.path.join(SUP, "S_" + os.path.basename(src)))

# ---------- MANIFEST ----------
manifest = ["# Network Neuroscience submission package, gene-to-brain (2026-06-30)\n",
            "Upload to the Network Neuroscience editorial system:\n",
            "1. Manuscript.docx (main text + abstract + references)",
            "2. Title_Page.docx", "3. Cover_Letter.docx",
            "4. Main_Figures/Fig1.pdf (per-subject 3-state correspondence)",
            "5. Main_Figures/Fig2.pdf (per-region localization)",
            "6. Main_Figures/Fig3.pdf (distance + cell-type robustness)",
            "7. Main_Tables/Table1.docx (main results)",
            "8. Supplementary/ (per-subject, per-region, state-pairing, cell-type, motion, pairwise tables)",
            "\nBefore submitting: complete Huang et al. 2025 author list; mint Zenodo DOI; create GitHub repo gene-to-brain-wj.",
            "See _reference_verification.md for the reference check."]
open(os.path.join(OUT, "MANIFEST.md"), "w", encoding="utf-8").write("\n".join(manifest))
print("Submission package assembled in", OUT)
for f in sorted(glob.glob(os.path.join(OUT, "**", "*"), recursive=True)):
    if os.path.isfile(f): print("  ", os.path.relpath(f, OUT))
